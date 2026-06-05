from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_PATH = Path(__file__).resolve().parent / "GyaanPlant_Admin_Dashboard_Test_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill="F2F4F7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level < 3 else 120)
    return paragraph


def add_metadata_table(doc):
    rows = [
        ("Application", "GyaanPlant LMS Admin Dashboard"),
        ("Environment URL", "https://lms.gyaanplant.co.in"),
        ("Automation Stack", "Python 3.12, Pytest 9.0.3, Playwright 1.60.0"),
        ("Design Pattern", "Page Object Model (POM)"),
        ("Execution Date", "May 31, 2026"),
        ("Execution Mode", "Headless Chromium"),
        ("Latest Result", "15 Passed, 0 Failed, 0 Skipped"),
        ("Duration", "61.59 seconds"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for field, value in rows:
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = value
    style_table(table)


def add_summary_table(doc):
    rows = [
        ("Smoke / Authentication", "3", "3", "0", "0"),
        ("Dashboard Happy Path", "5", "5", "0", "0"),
        ("Dashboard Feature Checks", "4", "4", "0", "0"),
        ("Smoke Click Crawl", "3", "3", "0", "0"),
        ("Total", "15", "15", "0", "0"),
    ]
    table = doc.add_table(rows=1, cols=5)
    headers = ["Suite", "Executed", "Passed", "Failed", "Skipped"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)


def add_test_case_table(doc):
    test_cases = [
        ("GP-SMOKE-001", "Authentication", "Dashboard route requires authentication", "Login form or login route is displayed", "PASSED"),
        ("GP-SMOKE-002", "Authentication", "Login page is ready for credentials", "Email input is visible and usable", "PASSED"),
        ("GP-SMOKE-003", "Authentication", "Admin credentials log in to dashboard", "Dashboard shell loads after login", "PASSED"),
        ("GP-ADM-DASH-001", "Sidebar Navigation", "Admin navigation modules are available", "All visible sidebar items render", "PASSED"),
        ("GP-ADM-DASH-002", "Top Metrics", "Dashboard KPI cards show expected values", "Partners 4, Students 25, Agreements 1, Employees 0, Activity 29", "PASSED"),
        ("GP-ADM-DASH-003", "Platform Activity", "Activity cards show expected values", "May 28 +785/29, May 29 +105/12, May 30 +50/10", "PASSED"),
        ("GP-ADM-DASH-004", "Top Partner Table", "Top partner table shows Microsoft", "Microsoft, IT, ACTIVE are visible", "PASSED"),
        ("GP-ADM-DASH-005", "Leaderboard", "Top five leaderboard entries are visible", "GYAANPLANT, ADA, RISHIK, PERSON1, SAGAR are visible", "PASSED"),
        ("GP-FEAT-001", "Sidebar Navigation", "Dashboard navigation menu displays admin modules", "Admin modules are visible", "PASSED"),
        ("GP-FEAT-002", "Summary Cards", "Platform metrics are visible", "Metric card labels are visible", "PASSED"),
        ("GP-FEAT-003", "Dashboard Widgets", "Activity and partner widgets are visible", "Activity summary and partner table headers render", "PASSED"),
        ("GP-FEAT-004", "Leaderboard Widget", "Leaderboard widget is visible", "Filters and full leaderboard action render", "PASSED"),
        ("GP-CLICK-001", "Scrollable Sidebar", "Click all sidebar pages including Finance and Platform", "Each route loads and returns to Dashboard", "PASSED"),
        ("GP-CLICK-002", "Leaderboard Filters", "Click Monthly and Weekly filters", "Leaderboard remains rendered", "PASSED"),
        ("GP-CLICK-003", "Header Actions", "Click Theme, Notifications, and Profile controls", "Page remains responsive and recovers to Dashboard", "PASSED"),
    ]
    table = doc.add_table(rows=1, cols=5)
    headers = ["Test ID", "Component", "Scenario", "Expected Result", "Status"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for case in test_cases:
        cells = table.add_row().cells
        for idx, value in enumerate(case):
            cells[idx].text = value
        set_cell_shading(cells[4], "E6F4EA")
    style_table(table)


def add_command_block(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "WEB_HEADLESS=true LMS_EMAIL=admin@gyaanplant.com LMS_PASSWORD=******** "
        ".venv/bin/python -m pytest tests/web -v"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.25)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("GyaanPlant Admin Dashboard Test Execution Report")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph("Functional happy path, smoke, and smoke-click automation coverage")
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "The GyaanPlant LMS admin dashboard web automation suite was executed successfully. "
        "The latest full web run completed with 15 passing tests and no failures. Coverage includes "
        "authentication readiness, dashboard happy path validation, exact dashboard data checks, "
        "scrollable sidebar navigation, leaderboard filters, and header controls."
    )
    add_metadata_table(doc)

    add_heading(doc, "2. Execution Summary", 1)
    add_summary_table(doc)

    add_heading(doc, "3. Executed Test Cases", 1)
    add_test_case_table(doc)

    add_heading(doc, "4. Test Command", 1)
    doc.add_paragraph("The following command was used for the latest full-suite execution:")
    add_command_block(doc)

    add_heading(doc, "5. Automation Artifacts", 1)
    for item in [
        "POM: pages/web/dashboard_page.py",
        "Base page helpers: pages/web/base_page.py",
        "Smoke tests: tests/web/test_dashboard_smoke.py",
        "Dashboard happy path tests: tests/web/test_admin_dashboard_happy_path.py",
        "Dashboard feature tests: tests/web/test_dashboard_features.py",
        "Smoke-click crawler: tests/web/test_dashboard_smoke_click.py",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "6. Notes", 1)
    for item in [
        "Credentials were supplied through environment variables and are not hard-coded in test files.",
        "The smoke-click crawler re-fetches locators inside each loop to avoid stale or detached element issues.",
        "Scrollable sidebar items are brought into view before clicking.",
        "Dashboard tests wait for real dashboard content to render after skeleton loaders disappear.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
    print(REPORT_PATH)
